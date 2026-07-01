#!/usr/bin/env python3
"""Create a DOCX report from the generated manuscript markdown."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
PAPER_DIR = ROOT / "paper"
INPUT = PAPER_DIR / "manuscript_results_filled.md"
OUTPUT = PAPER_DIR / "manuscript_results_filled.docx"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "D0D7DE")
        borders.append(tag)
    tbl_pr.append(borders)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.tcW
            tc_w.set(qn("w:w"), str(widths_dxa[min(idx, len(widths_dxa) - 1)]))
            tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGIN_DXA.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def column_widths(max_cols: int) -> list[int]:
    if max_cols >= 9:
        weights = [1.25, 1.2, 0.55, 0.75, 1.05, 1.05, 1.1, 1.0, 1.0][:max_cols]
    elif max_cols == 7:
        weights = [1.8, 1.2, 0.65, 0.8, 1.0, 1.0, 1.1]
    elif max_cols == 5:
        weights = [1.9, 1.3, 0.7, 0.9, 1.1]
    else:
        weights = [1.0] * max_cols
    total = sum(weights)
    raw = [round(TABLE_WIDTH_DXA * w / total) for w in weights]
    raw[-1] += TABLE_WIDTH_DXA - sum(raw)
    return raw


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.text = "Updated Health-AI Readiness Public-Subset Replication"
    header.paragraph_format.space_after = Pt(0)
    for run in header.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("666666")

    footer = section.footer.paragraphs[0]
    footer.text = "Generated from local benchmark outputs"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("666666")

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_inline_markup(paragraph, text: str) -> None:
    text = text.replace("`", "")
    pos = 0
    for match in re.finditer(r"\*\*(.*?)\*\*", text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        run = paragraph.add_run(match.group(1))
        run.bold = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        parts = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if parts and not all(set(cell) <= {"-", ":"} for cell in parts):
            rows.append(parts)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    set_table_geometry(table, column_widths(max_cols))
    set_table_borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx in range(max_cols):
            cell = table.cell(r_idx, c_idx)
            set_cell_margins(cell)
            text = row[c_idx] if c_idx < len(row) else ""
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.size = Pt(8 if max_cols >= 8 else 9)
                    run.font.name = "Calibri"
            if r_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    doc.add_paragraph()


def resolve_image(path_text: str) -> Path:
    raw = Path(path_text)
    if raw.is_absolute():
        return raw
    return (PAPER_DIR / raw).resolve()


def build_docx() -> None:
    if not INPUT.exists():
        raise SystemExit(f"Missing {INPUT}. Run build_paper_from_results.py first.")

    doc = Document()
    style_document(doc)
    lines = INPUT.read_text().splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue

        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, parse_table(table_lines))
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            alt, path_text = image_match.groups()
            image_path = resolve_image(path_text)
            if image_path.exists():
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(str(image_path), width=Inches(6.0))
                caption = doc.add_paragraph(alt)
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in caption.runs:
                    run.italic = True
                    run.font.size = Pt(9)
            i += 1
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor.from_string("0B2545")
            i += 1
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
            i += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            i += 1
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markup(p, line[2:].strip())
            i += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.*)", line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_inline_markup(p, numbered.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_markup(p, line)
        i += 1

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build_docx()
